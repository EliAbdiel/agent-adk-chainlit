import fitz
import base64
import asyncio
import aiofiles
import subprocess
import chainlit as cl
from io import BytesIO
from google import genai
from pathlib import Path
from docx import Document
from google.genai import types
from typing import Optional, Set
from src.logs.logger import setup_logger
from typing import Optional, Dict, Any, List
from src.document.processor_config import ProcessingConfig
from src.config.common import GEMINI_API_KEY

class DocumentProcessor:
    """
    Enhanced document processor with better error handling,
    support for more file types, and advanced features.
    """
    def __init__(
            self, 
            gemini_api_key: str = GEMINI_API_KEY, 
            config: Optional[ProcessingConfig] = None,
        ) -> None:
        self.config = config or ProcessingConfig()
        self.client = genai.Client(api_key=gemini_api_key)
        self.logger = setup_logger('DOCS PROCESSOR')
        self.file_processor_map = {
            '.pdf': self._extract_text_from_pdf_bytes,
            '.docx': self._extract_text_from_docx_bytes,
            '.txt': self._extract_text_from_txt_bytes,
            '.jpg': self._extract_content_from_image_bytes,
            '.jpeg': self._extract_content_from_image_bytes,
            '.png': self._extract_content_from_image_bytes,
        }

    def _read_bytes(self, file: cl.File) -> bytes:
        """Read file content as bytes, handling both in-memory and disk storage"""
        content = getattr(file, 'content', None)
        if isinstance(content, (bytes, bytearray)):
            return bytes(content)
        file_path = getattr(file, 'path', None)
        if isinstance(file_path, str) and file_path:
            with open(file_path, 'rb') as f:
                return f.read()
        raise ValueError("File content is not available as bytes or valid path")

    async def _get_expected_mime_types(self, extension: str) -> Set[str]:
        """Get expected MIME types for a given file extension"""
        return self.config.allowed_mime_types.get(extension.lower(), set())

    async def _get_file_info(self, filename: str, file_bytes: bytes, file_mime: str) -> Dict[str, Any]:
        """Extract file information and validate"""
        file_extension = Path(filename).suffix.lower()
        file_size = len(file_bytes)

        return {
            'filename': filename,
            'extension': file_extension,
            'size': file_size,
            'mime_type': file_mime
        }
    
    async def _validate_file(self, filename: str, file_bytes: bytes, file_mime: str) -> Dict[str, Any]:
        """Enhanced file validation with content type checking"""
        file_info = await self._get_file_info(filename=filename, file_bytes=file_bytes, file_mime=file_mime)

        # Validate file extension
        if file_info['extension'] not in self.config.allowed_extensions:
            raise ValueError(f"Unsupported file extension: {file_info['extension']}")

        # Validate file content matches extension
        expected_mimes = await self._get_expected_mime_types(extension=file_info['extension'])
        if expected_mimes and file_info['mime_type'] not in expected_mimes:
            raise ValueError(
                f"File content doesn't match extension: {file_info['extension']}. "
                f"Expected: {expected_mimes}, Got: {file_info['mime_type']}"
            )

        # Check file size
        if file_info['size'] > self.config.max_file_size:
            raise ValueError(f"File size ({file_info['size']} bytes) exceeds limit ({self.config.max_file_size} bytes)")

        # Validate filename for security
        if '..' in filename or filename.startswith('/'):
            raise ValueError("Invalid filename: potential path traversal detected")

        return file_info

    async def _clean_and_summarize_text(self, text: str, filename: str, doc_type: str = "document") -> str:
        """
        Clean and summarize text extracted from pdf, docx, txt, or images.
        Enhanced with better prompts and error handling.
        """
        if not text or not text.strip():
            return "No extractable text content found."

        try:
            # Enhanced prompt for better summarization
            system_instruction = f"""You are a summarization assistant specialized in long-form content.

            Input context
            - **Filename**: "{filename}"
            - **Document Category**: "{doc_type}"

            When summarizing, you must:
            - Preserve original meaning, intent, and logical flow
            - Extract only the most relevant information
            - Remove redundancy while keeping factual accuracy
            - Use clear, structured, concise language

            Constraints:
            - No opinions, assumptions, or invented facts
            - No altering context beyond compression
            - If ambiguous, summarize only what is certain

            Tone & Style:
            - Neutral, professional, short, clear, direct
            - Summary length: 10–30% of original
            - Always end with a 1-sentence summary (≤20 words)

            Mandatory Output Format:
            1. Title (only if original had one)
            *Short, accurate summary title*

            2. Executive Summary (1 paragraph, 3–5 sentences)
            *Concise overview of full content*

            3. Main Points (bullet list, 3+ key ideas)
            - Key idea 1
            - Key idea 2
            - Key idea 3
            - Additional essential details if needed

            4. Section Breakdown (only if original had multiple topics)
            **Section A — Topic**
            - Highlight 1
            - Highlight 2

            **Section B — Topic**
            - Highlight 1
            - Highlight 2

            5. Important Data & Facts (only if useful)
            | Fact/Metric | Detail |
            |---|---|
            | Example | Result |

            6. Key Takeaways (3–5 insights)
            ✅ Insight 1  
            ✅ Insight 2  
            ✅ Insight 3  

            7. One-Sentence Summary (≤20 words)
            *Factual compression of entire content*

            8. Tags (only if original topics are identifiable)
            topic1, topic2

            Final Rules:
            - Follow the format exactly
            - Always include sections 2, 3, 6, 7
            - Exclude optional sections if they add no value
            - Make the summary scannable and fact-driven"""

            response = self.client.models.generate_content(
                model=self.config.gemini_model,
                config=types.GenerateContentConfig(
                    system_instruction=system_instruction,   
                    temperature=self.config.temperature,
                ),
                contents=text,
            )

            self.logger.info(f"Successfully processed {doc_type} text with Gemini API")
            return response.text

        except Exception as e:
            self.logger.error(f"Error in text cleaning/summarization: {str(e)}")
            # Return original text if summarization fails
            return f"Original content:\n\n{text[:self.config.text_extract_limit]}"

    async def _extract_text_from_pdf_bytes(self, filename: str, pdf_bytes: bytes, pdf_mime: str) -> str:
        """
        Extract text from PDF bytes using PyMuPDF with enhanced features.
        """
        try:

            doc = fitz.open(stream=pdf_bytes, filetype='pdf')
            text = ''
            
            self.logger.info(f"Processing PDF with {len(doc)} pages")

            # Extract text from each page
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text:
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"

            doc.close()

            # Log extraction summary
            self.logger.info(f"Extracted {len(text)} characters from PDF")

            result = await self._clean_and_summarize_text(text=text[:self.config.text_extract_limit], filename=filename, doc_type=pdf_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing PDF: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")

    async def _extract_text_from_docx_bytes(self, filename: str, docx_bytes: bytes, docx_mime: str) -> str:
        """
        Extract text from DOCX bytes with enhanced features.
        """
        try:

            doc = Document(BytesIO(docx_bytes))
            text = ''

            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells if cell.text)
                    if row_text:
                        text += f"Table row: {row_text}\n"

            self.logger.info(f"Extracted {len(text)} characters from DOCX")

            result = await self._clean_and_summarize_text(text=text[:self.config.text_extract_limit].strip(), filename=filename, doc_type=docx_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing DOCX: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")

    async def _extract_text_from_txt_bytes(self, filename: str, txt_bytes: bytes, txt_mime= str) -> str:
        """
        Extract text from plain text files.
        """
        try:

            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1']
            text = None

            for encoding in encodings:
                try:
                    text = txt_bytes.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if text is None:
                # Fallback: replace errors
                text = txt_bytes.decode('utf-8', errors='replace')

            self.logger.info(f"Extracted {len(text)} characters from TXT")

            result = await self._clean_and_summarize_text(text=text[:self.config.text_extract_limit], filename=filename, doc_type=txt_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing TXT: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")

    async def _extract_content_from_image_bytes(self, filename: str, image_bytes: bytes, image_mime: str) -> str:
        """
        Extract content from image bytes using gemini as a vision model.
        """
        try:

            # Enhanced prompt for better extraction
            system_instruction = """You are a highly accurate OCR and document extraction engine.
            Your task is to convert the provided image into text while preserving the original structure and formatting as closely as possible.

            Follow these rules:
            1. **Text Fidelity:** Transcribe text exactly as it appears. Do not summarize or correct grammar.
            2. **Tables:** If the image contains tables, represent them using Markdown table syntax.
            3. **Headings:** Use Markdown headers (#, ##, ###) to represent titles and section headings.
            4. **Lists:** Use Markdown bullet points or numbered lists to represent list items.
            5. **Legibility:** If parts of the text are illegible or cut off, output "[Illegible]" or "[Cut off]" for those specific sections.
            6. **No Conversational Filler:** Do not start with "Here is the text" or "Sure." Output *only* the extracted content.
            """

            response = self.client.models.generate_content(
                model=self.config.gemini_model,
                contents=[
                    types.Part.from_bytes(
                        data=image_bytes,
                        mime_type=image_mime,
                    ),
                    system_instruction
                ]
            )

            self.logger.info("Successfully processed image with multimodal model")

            result = await self._clean_and_summarize_text(text=response.text, filename=filename, doc_type=image_mime)

            return result

        except Exception as e:
            self.logger.error(f"Error processing image: {str(e)}")
            raise ValueError(f"Failed to process image: {str(e)}")
            
    @cl.step(name="process-document", type="tool", show_input=False)
    async def process_document_async(self, filename: str, file_bytes: bytes, file_mime: str) -> str:
        """
        Process a document and extract structured content.

        Args:
            filename (str): The name of the file
            file_bytes (bytes): The raw bytes of the file
            file_mime (str): The mime type of the file

        Returns:
            str: Processed and summarized content
        """
        try:
            # Validate input
            if not filename or not file_bytes:
                raise ValueError("Filename and file_bytes are required")

            # Validate file and get info
            file_info = await self._validate_file(filename=filename, file_bytes=file_bytes, file_mime=file_mime)
            self.logger.info(f"Processing file: {file_info['filename']} "
                           f"({file_info['size']} bytes)")

            # Get appropriate processor
            processor = self.file_processor_map.get(file_info['extension'])
            if not processor:
                raise ValueError(f"Unsupported file extension: {file_info['extension']}")

            # Process the document
            result = await processor(filename, file_bytes, file_mime)

            self.logger.info(f"Successfully processed {filename}")
            return result

        except Exception as e:
            self.logger.error(f"Failed to process {filename}: {str(e)}")
            raise

    def batch_process_documents(self, file_data: Dict[str, bytes]) -> Dict[str, str]:
        """
        Process multiple documents in a batch.

        Args:
            file_data (Dict[str, bytes]): Dictionary of filename -> file_bytes

        Returns:
            Dict[str, str]: Dictionary of filename -> processed content
        """
        results = {}

        for filename, file_bytes in file_data.items():
            try:
                results[filename] = self.process_document_async(filename, file_bytes)
            except Exception as e:
                results[filename] = f"Error processing {filename}: {str(e)}"

        return results

    async def batch_process_documents_async(self, file_data: Dict[str, bytes]) -> Dict[str, str]:
        """
        Process multiple documents in a batch asynchronously.

        Args:
            file_data (Dict[str, bytes]): Dictionary of filename -> file_bytes

        Returns:
            Dict[str, str]: Dictionary of filename -> processed content
        """
        results = {}

        # Process files concurrently
        tasks = []
        for filename, file_bytes in file_data.items():
            task = asyncio.create_task(self.process_single_file_async(filename, file_bytes))
            tasks.append((filename, task))

        for filename, task in tasks:
            try:
                results[filename] = await task
            except Exception as e:
                results[filename] = f"Error processing {filename}: {str(e)}"

        return results

    async def process_single_file_async(self, file: cl.File) -> str:
        """Helper method for async processing of a single file"""
        if file is None:
            raise ValueError("File is None")
        filename = str(file.name)
        file_bytes = self._read_bytes(file=file)
        file_mime = str(file.mime)
        
        self.logger.info(f"Processing single file: {filename}, size: {len(file_bytes)} bytes, mime: {file_mime}")
        
        content = await self.process_document_async(filename=filename, file_bytes=file_bytes, file_mime=file_mime)
        
        return content
    
    @cl.step(name="summarize-content", type="tool", show_input=False)
    async def summarize_text(self, content: str) -> str:
        """
        Summarize text using the document processor.

        Args:
            content (str): The content to summarize

        Returns:
            str: Summarized content
        """
        return await self._clean_and_summarize_text(text=content, filename="Document", doc_type="document")
